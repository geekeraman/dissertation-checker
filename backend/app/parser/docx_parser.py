"""DOCX document parser — extracts structured data from .docx files."""

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from app.parser.structures import (
    ParsedDocument, ParsedParagraph, DocumentSection,
    Figure, Table, Reference, DocumentMetadata, DocProperties,
)
import re


# Ключевые слова для распознавания структурных заголовков по тексту
_STRUCTURAL_HEADING_KEYWORDS = {
    "аңдатпа", "реферат", "аннотация", "abstract",
    "мазмұны", "мазмұн", "содержание", "оглавление", "contents",
    "кіріспе", "введение", "introduction",
    "негізгі бөлім", "основная часть",
    "қорытынды", "заключение", "conclusion",
    "пайдаланылған дереккөздер тізімі", "әдебиеттер тізімі",
    "список литературы", "список использованных источников",
    "references", "bibliography",
    "қосымшалар", "қосымша", "приложения", "приложение",
    "appendix", "appendices",
}


def _is_structural_heading_text(text: str) -> bool:
    """Проверяет, совпадает ли текст с известными структурными заголовками."""
    stripped = text.strip()
    # Реальные заголовки короткие — длинный текст точно не заголовок
    if len(stripped) > 100:
        return False
    lower = stripped.lower()
    # Убираем числовую нумерацию вначале (напр. '1 КІРІСПЕ' -> 'кіріспе')
    clean = re.sub(r'^\d+(\.\d+)*[.\s]*', '', lower).strip()
    for kw in _STRUCTURAL_HEADING_KEYWORDS:
        if lower == kw or clean == kw:
            return True
    return False


ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    None: None,
}


def _extract_font_info(paragraph) -> tuple[str | None, float | None, bool | None]:
    """Extract font name, size, and bold from a paragraph's runs."""
    font_name = None
    font_size = None
    bold = None
    for run in paragraph.runs:
        if run.font.name and font_name is None:
            font_name = run.font.name
        if run.font.size and font_size is None:
            font_size = run.font.size.pt
        if run.font.bold is not None and bold is None:
            bold = run.font.bold
    return font_name, font_size, bold


def _detect_figures(doc: Document, paragraphs: list[ParsedParagraph]) -> list[Figure]:
    """Detect figures from document content and inline shapes."""
    figures = []
    figure_pattern = re.compile(r"[Сс]урет\s*(\d+(?:\.\d+)?)", re.UNICODE)

    for i, para in enumerate(paragraphs):
        match = figure_pattern.search(para.text)
        if match:
            has_image = False
            # Check for inline shapes in the original docx paragraph
            if i < len(doc.paragraphs):
                for run in doc.paragraphs[i].runs:
                    if run._element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
                    ) or run._element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'
                    ):
                        has_image = True
                        break

            figures.append(Figure(
                number=match.group(1),
                title=para.text,
                paragraph_index=i,
                has_caption=True,
                caption_paragraph_index=i,
                caption_position="below",
            ))
    return figures


def _detect_tables(doc: Document, paragraphs: list[ParsedParagraph]) -> list[Table]:
    """Detect tables and their captions from the document."""
    tables = []
    table_pattern = re.compile(r"(?:Кесте|Таблица|Table)\s*(\d+(?:\.\d+)?)", re.UNICODE)

    for i, para in enumerate(paragraphs):
        match = table_pattern.search(para.text)
        if match:
            tables.append(Table(
                number=match.group(1),
                title=para.text,
                paragraph_index=i,
                has_caption=True,
                caption_paragraph_index=i,
                caption_position="above",
            ))
    return tables


def _detect_references(paragraphs: list[ParsedParagraph]) -> list[Reference]:
    """Detect reference entries after the references heading."""
    references = []
    in_references = False
    ref_headings = {"тізімі", "список", "references", "әдебиеттер"}

    for i, para in enumerate(paragraphs):
        if para.is_heading:
            lower = para.text.lower()
            if any(rh in lower for rh in ref_headings):
                in_references = True
                continue
            elif in_references and para.heading_level == 1:
                break

        if in_references and para.text.strip():
            references.append(Reference(text=para.text, paragraph_index=i))

    return references


def _build_sections(paragraphs: list[ParsedParagraph]) -> list[DocumentSection]:
    """Build document sections from headings.

    Распознаёт разделы по двум критериям:
    1. Параграф имеет стиль 'Heading 1'
    2. Текст параграфа совпадает с известным структурным заголовком
       и параграф отформатирован как заголовок (жирный, по центру или прописные буквы)
    """
    sections = []
    for i, para in enumerate(paragraphs):
        is_section_heading = False

        if para.is_heading and para.heading_level == 1:
            is_section_heading = True
        elif _is_structural_heading_text(para.text):
            # Параграф не помечен стилем Heading, но текст совпадает
            # с известным структурным заголовком — проверяем форматирование
            text = para.text.strip()
            is_bold = para.bold is True
            is_centered = para.alignment == "center"
            is_upper = text == text.upper() and len(text) > 2
            if is_bold or is_centered or is_upper:
                is_section_heading = True

        if is_section_heading:
            sections.append(DocumentSection(
                name=para.text.strip(),
                heading=para.text.strip(),
                start_paragraph_index=i,
                end_paragraph_index=len(paragraphs) - 1,
                level=1,
            ))
    # Set end_paragraph_index for each section
    for j in range(len(sections) - 1):
        sections[j].end_paragraph_index = sections[j + 1].start_paragraph_index - 1
    return sections


def _extract_properties(doc: Document) -> DocProperties:
    """Extract document-level properties from section settings."""
    props = DocProperties()
    section = doc.sections[0] if doc.sections else None
    if section:
        if section.left_margin:
            props.left_margin_cm = section.left_margin.cm
        if section.right_margin:
            props.right_margin_cm = section.right_margin.cm
        if section.top_margin:
            props.top_margin_cm = section.top_margin.cm
        if section.bottom_margin:
            props.bottom_margin_cm = section.bottom_margin.cm
        if section.page_width:
            props.page_width_cm = section.page_width.cm
        if section.page_height:
            props.page_height_cm = section.page_height.cm

    # Try to get default font from styles
    default_style = doc.styles["Normal"] if "Normal" in [s.name for s in doc.styles] else None
    if default_style and default_style.font:
        props.default_font_name = default_style.font.name
        if default_style.font.size:
            props.default_font_size = default_style.font.size.pt

    return props


def _estimate_page_count(paragraphs: list[ParsedParagraph]) -> int:
    """Rough page count estimation based on paragraph content."""
    total_chars = sum(len(p.text) for p in paragraphs)
    # A4 with TNR 14pt 1.5 spacing ≈ 1800 chars per page
    return max(1, total_chars // 1800)


def parse_docx(file_path: str, doc_type: str) -> ParsedDocument:
    """Parse a .docx file and return a ParsedDocument."""
    doc = Document(file_path)
    paragraphs: list[ParsedParagraph] = []

    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Normal"
        is_heading = style_name.startswith("Heading") or "heading" in style_name.lower()
        heading_level = None
        if is_heading:
            try:
                heading_level = int(style_name.replace("Heading", "").replace("heading", "").strip())
            except (ValueError, TypeError):
                heading_level = 1

        font_name, font_size, bold = _extract_font_info(para)
        alignment = ALIGNMENT_MAP.get(para.alignment)

        line_spacing = None
        if para.paragraph_format.line_spacing is not None:
            try:
                spacing_rule = para.paragraph_format.line_spacing_rule
                raw_value = float(para.paragraph_format.line_spacing)
                # Если правило — MULTIPLE (или None/по умолчанию), значение — множитель
                if spacing_rule is None or spacing_rule == WD_LINE_SPACING.MULTIPLE:
                    line_spacing = raw_value
                else:
                    # EXACTLY / AT_LEAST — значение в Twips (1/20 pt)
                    # Если значение > 10, это точно не множитель — пропускаем
                    if raw_value > 10:
                        line_spacing = None
                    else:
                        line_spacing = raw_value
            except (ValueError, TypeError):
                pass

        first_line_indent = None
        if para.paragraph_format.first_line_indent:
            try:
                first_line_indent = para.paragraph_format.first_line_indent.cm
            except (ValueError, TypeError):
                pass

        has_page_break = False
        for run in para.runs:
            if run._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type="page"]'
            ):
                has_page_break = True
                break

        paragraphs.append(ParsedParagraph(
            text=para.text,
            style_name=style_name,
            is_heading=is_heading,
            heading_level=heading_level,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            alignment=alignment,
            line_spacing=line_spacing,
            first_line_indent=first_line_indent,
            has_page_break_before=has_page_break,
            paragraph_index=idx,
        ))

    sections = _build_sections(paragraphs)
    figures = _detect_figures(doc, paragraphs)
    tables = _detect_tables(doc, paragraphs)
    references = _detect_references(paragraphs)
    properties = _extract_properties(doc)
    page_count = _estimate_page_count(paragraphs)

    return ParsedDocument(
        doc_type=doc_type,
        paragraphs=paragraphs,
        sections=sections,
        figures=figures,
        tables=tables,
        references=references,
        metadata=DocumentMetadata(
            title=doc.core_properties.title,
            author=doc.core_properties.author,
        ),
        page_count=page_count,
        page_count_body=page_count,
        properties=properties,
    )
