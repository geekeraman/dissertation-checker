"""Tests for the DOCX parser."""

import pytest
from unittest.mock import patch, MagicMock
from app.parser.docx_parser import parse_docx
from app.parser.structures import ParsedDocument


class TestParseDocx:
    def test_returns_parsed_document_with_correct_doc_type(self, tmp_path):
        # Create a minimal .docx file for testing
        doc_path = tmp_path / "test.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test paragraph")
        doc.save(str(doc_path))

        result = parse_docx(str(doc_path), "thesis_science")
        assert isinstance(result, ParsedDocument)
        assert result.doc_type == "thesis_science"

    def test_extracts_paragraphs(self, tmp_path):
        doc_path = tmp_path / "test.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.save(str(doc_path))

        result = parse_docx(str(doc_path), "project")
        assert len(result.paragraphs) >= 2
        assert result.paragraphs[0].text == "First paragraph"

    def test_extracts_headings(self, tmp_path):
        doc_path = tmp_path / "test.docx"
        from docx import Document
        doc = Document()
        doc.add_heading("КІРІСПЕ", level=1)
        doc.add_paragraph("Body text")
        doc.save(str(doc_path))

        result = parse_docx(str(doc_path), "thesis_humanities")
        headings = [p for p in result.paragraphs if p.is_heading]
        assert len(headings) >= 1
        assert headings[0].heading_level == 1

    def test_extracts_document_properties(self, tmp_path):
        doc_path = tmp_path / "test.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(str(doc_path))

        result = parse_docx(str(doc_path), "project")
        assert result.properties is not None

    def test_extracts_sections(self, tmp_path):
        doc_path = tmp_path / "test.docx"
        from docx import Document
        doc = Document()
        doc.add_heading("КІРІСПЕ", level=1)
        doc.add_paragraph("Introduction text")
        doc.add_heading("1 НЕГІЗГІ БӨЛІМ", level=1)
        doc.add_paragraph("Main body text")
        doc.save(str(doc_path))

        result = parse_docx(str(doc_path), "thesis_science")
        assert len(result.sections) >= 1
