"""End-to-end integration tests."""

import pytest
from fastapi.testclient import TestClient
from app.main import app
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


@pytest.fixture
def client():
    return TestClient(app)


def create_sample_docx() -> io.BytesIO:
    """Create a sample .docx with known issues for integration testing."""
    doc = Document()

    # Set correct margins
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Title page
    doc.add_paragraph("ДИССЕРТАЦИЯ")

    # Missing Abstract section (should be flagged)
    # Contents
    doc.add_heading("МАЗМҰНЫ", level=1)
    doc.add_paragraph("1 Кіріспе .......... 3")

    # Introduction
    doc.add_heading("КІРІСПЕ", level=1)
    doc.add_paragraph("This is the introduction text with  double spaces.")

    # Main body
    doc.add_heading("1 НЕГІЗГІ БӨЛІМ", level=1)
    para = doc.add_paragraph("Body text here.")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    # Conclusion
    doc.add_heading("ҚОРЫТЫНДЫ", level=1)
    doc.add_paragraph("Conclusion text.")

    # References
    doc.add_heading("ӘДЕБИЕТТЕР ТІЗІМІ", level=1)
    doc.add_paragraph("[1] Author, Title, 2024")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


class TestIntegration:
    def test_health_endpoint(self, client):
        response = client.get("/api/health")
        assert response.status_code == 200
        assert response.json() == {"status": "ok"}

    def test_check_endpoint_returns_report(self, client):
        docx_file = create_sample_docx()
        response = client.post(
            "/api/check",
            files={"file": ("test.docx", docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"doc_type": "thesis_science"},
        )
        assert response.status_code == 200
        data = response.json()
        assert "id" in data
        assert "issues" in data
        assert data["filename"] == "test.docx"
        assert data["doc_type"] == "thesis_science"
        assert data["total_issues"] > 0

    def test_check_rejects_non_docx(self, client):
        response = client.post(
            "/api/check",
            files={"file": ("test.pdf", b"not a docx", "application/pdf")},
            data={"doc_type": "thesis_science"},
        )
        assert response.status_code == 400

    def test_report_contains_expected_categories(self, client):
        docx_file = create_sample_docx()
        response = client.post(
            "/api/check",
            files={"file": ("test.docx", docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"doc_type": "thesis_science"},
        )
        data = response.json()
        categories = set(data["issues_by_category"].keys())
        # At least structure issues should be present (missing abstract)
        assert "structure" in categories
